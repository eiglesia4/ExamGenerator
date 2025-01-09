# This is a sample Python script.

# Press ⌃R to execute it or replace it with your code.
# Press Double ⇧ to search everywhere for classes, files, tool windows, actions, and settings.

import pandas as pd
import argparse
import random
import numpy as np
from docx import Document

def     create_exam(df, df_long, long_titles, answer_columns, question='Pregunta', description=None, subject=None, is_odd=False):
    document = Document()
    if subject is None:
        subject = 'Examen Psicobiología'
    p = document.add_paragraph(f'{subject}', style='Heading 1')
    p.alignment=1
    document.add_paragraph('Apellidos y Nombre:', style='Normal')
    document.add_paragraph('Curso:', style='Normal')

    if description is not None:
        description_file = open(description, 'r')
        # for line in description_file: check if line starts wit tab, count tab and add paragraph with style List Bullet
        for line in description_file:
            if line.startswith('\t'):
                tabs = line.count('\t')
                if tabs == 1:
                    p = document.add_paragraph(f'{line.strip()}', style='List Bullet')
                else:
                    p = document.add_paragraph(f'{line.strip()}', style=f'List Bullet {tabs}')
            else:
                p = document.add_paragraph(f'{line.strip()}', style=f'Normal')
    else :
        p0 = document.add_paragraph(f'5 puntos en 20 preguntas tipo test (o respuesta en una o dos palabras): coeficiente de reducción cada 3 resta 0,25', style='List Bullet')
        p1 = document.add_paragraph(f'5 puntos en preguntas de desarrollo:', style='List Bullet')
        p2 = document.add_paragraph(f'Un problema de genética.', style='List Bullet 2')
        p3 = document.add_paragraph(f'3 preguntas del temario.', style='List Bullet 2')
        p4 = document.add_paragraph(f'Una pregunta sobre el trabajo final: debe ir en la línea del trabajo entregado, será la fuente principal para la corrección de este apartado. Puntos a tratar:', style='List Bullet 2')
        p5 = document.add_paragraph(f'Tema tratado, ¿cómo afecta a la conducta?', style='List Bullet 3')
        p6 = document.add_paragraph(f'Causas (genéticas, ambientales, de los dos tipos...)', style='List Bullet 3')
        p7 = document.add_paragraph(f'Abordaje desde el punto de vista profesional.', style='List Bullet 3')

    document.add_paragraph(f'')
    document.add_paragraph(f'El examen durará aproximadamente 1:30', style='Normal')
    document.add_paragraph(f'')

    # PARTE TIPO TEST
    question_num = 1
    _num_id = document.get_new_list("10")
    for index, row in df.iterrows():
        q = row[question]
        question_num = question_num + 1
        if not q.endswith(""):
            q = f'{q}:'
        p = document.add_paragraph(f'{q}', style='ListParagraph')
        p.num_id = _num_id
        p.level = 0
        # document.add_paragraph(f'{index+1}.- {q}')
        a = 1
        for col in answer_columns:
            ans = f'{df.loc[index][col]}'
            if ans != "nan":
                if ans == 'void':
                    ans = ''
                    p = document.add_paragraph(f'{ans}', style='Normal')
                else:
                    p = document.add_paragraph(f'{ans}', style='ListParagraph')
                    p.num_id = _num_id
                    p.level = 1
                    a = a+1
        document.add_paragraph(f'')


    # PARTE DESARROLLO
    for _i in range(len(df_long)):
        _num_id = document.get_new_list("10")
        document.add_paragraph(f'')
        document.add_paragraph(f'{long_titles[_i]}', style='Normal')
        for index, row in df_long[_i].iterrows():
            q = row[question]
            if not q.endswith(""):
                q = f'{q}:'
            p = document.add_paragraph(f'{q}', style='ListParagraph')
            p.num_id = _num_id
            p.level = 0
        document.add_paragraph(f'')

    document.add_paragraph(f'')
    return document


def print_exam(df,  answer_columns, question='Pregunta'):
    for index, row in df.iterrows():
        q = row[question]
        if not q.endswith(""):
            q = f'{q}:'
        print(f'{index+1}.- {q}')
        a = 1
        for col in answer_columns:
            ans = f'{df.loc[index][col]}'
            if ans != "nan":
                print(f'\t{a}.- {ans}')
                a = a+1
    print()

# Press the green button in the gutter to run the script.
if __name__ == '__main__':
    """
    Generates a number of exams from a csv file with the format Question;answer1;answer2; ...;answerN
    Also a number of long questions csv files can be added with the format Question\nQuestion 1\nQuestion 2\n...\nQuestion N
    You can add several long questions files and the number of questions to be used for each file
    Also you can provide titles for each long question file
    BUT the number of long questions, number of questions and titles must be the same
    
    example: python3 -u main_docx.py -f examen_23_test.csv -s "Examen de Psicobiología" -d description.txt -e 4 -p examen2023_ -n 20 -l examen_23_problema_genetica.csv examen_23_preguntas_temario.csv examen_23_perguntas_trabajo.csv -x 1 3 1 -t "Problema de genética:" "Preguntas de temario:" "Pregunta sobre el trabajo final:"
    """
    parser = argparse.ArgumentParser(
                    description = 'Takes a csv with an exam in format Question;answer1;answer2; ...;answerN and prints a shuffled version',
                    epilog = 'Warning all columns as shuffled equally')
    parser.add_argument('-f', '--file', required=True, help="The file with the exam to shuffle")
    parser.add_argument('-e', '--exams', required=True, type=int,  help="The number of generated documents file with the exam to shuffle")
    parser.add_argument('-s', '--subject', required=False, type=str,  help="The name of the subject. If not defined the subject is Examen de Psicobiologia")
    parser.add_argument('-d', '--description', required=False, type=str,  help="A file with the description of the exam. If not defined the description is empty")
    parser.add_argument('-p', '--prefix', nargs='?', default='exam', type=str, help="The prefix for the output filename")
    parser.add_argument('-n', '--number', required=False, type=int,  help="The number of questions to insert if not defined all questions are used (optional)")
    parser.add_argument('-l', '--long', required=False, nargs='+', type=str,  help="Longquestions. The list of csv for long questions (optional). The first line of each file must be question")
    parser.add_argument('-x', '--numslong', required=False, nargs='+', type=int,  help="Long Questions. The list of number of long questions (optional)")
    parser.add_argument('-t', '--titleslong', required=False, nargs='+', type=str,  help="Long Questions. The list of titles for long questions (optional)")
    args = vars(parser.parse_args())
    df = None
    df_opt2 = []
    opt2_titles = []

    # CHECK LONG QUESTIONS
    if args['long'] is not None:
        longquestions = args['long']
        numslong = args['numslong']
        titleslong = args['titleslong']
        if len(longquestions) != len(numslong) or len(longquestions) != len(titleslong):
            print("Error: You must specify the same number of long questions, number of questions and titles")
            exit(1)
        opt2_titles = args['titleslong']


    # GENERATE EXAMS
    for i in range(args['exams']):
        print(f'Generating exam {i}')
        # PROCESS TEST QUESTIONS
        df = pd.read_csv(args['file'], delimiter=';', index_col=False)
        df = df.reindex(np.random.permutation(df.index))
        df = df.reset_index(drop=True)
        numquestions = len(df.index)
        if args['number'] is not None:
            numquestions = args['number']
        df = df.head(numquestions)

        # PROCESS LONG QUESTIONS
        df_opt2 = []
        if args['long'] is not None:
            for j in range(len(longquestions)):
                df_long = pd.read_csv(longquestions[j], delimiter=';', index_col=False)
                df_long = df_long.reindex(np.random.permutation(df_long.index))
                df_long = df_long.reset_index(drop=True)
                df_long = df_long.reindex(np.random.permutation(df_long.index))
                df_long = df_long.reset_index(drop=True)
                df_long = df_long.reindex(np.random.permutation(df_long.index))
                df_long = df_long.reset_index(drop=True)
                df_long = df_long.reindex(np.random.permutation(df_long.index))
                df_long = df_long.reset_index(drop=True)
                df_long = df_long.head(numslong[j])
                df_opt2.append(df_long)

        doc = create_exam(df=df, subject=args["subject"], description=args["description"], df_long=df_opt2, long_titles=opt2_titles, answer_columns=df.columns.values.tolist()[1:], question=df.columns[0], is_odd=i%2==0)
        file_name = f'{args["prefix"]}{i}.docx'
        doc.save(file_name)
        print(f'Saved file {file_name}')

    # question = df.columns[0]
    # _columns = df.columns.values.tolist()[1:]
    # #    print_exam(df, _columns, question)
    # numfirst = len(df.index) - numsecond
    # if args['number'] is not None:
    #     numfirst = args['number'] - numsecond
    #
    # df_long = None
    # if args['longquestions'] is not None:
    #     df_long = pd.read_csv(args['longquestions'], delimiter=';', index_col=False)
    #
    # for i in range(args['exams']):
    #     random.shuffle(_columns)
    #     df1 = df.reindex(np.random.permutation(df.index))
    #     df2 = df1.reset_index(drop=True)
    #     df2 = df2.head(numfirst)
    #     print(df2)
    #     if df_opt2 is not None:
    #         df2 = pd.concat([df2, df_opt2], ignore_index=True)
    #     if df_long is not None:
    #         df_long = df_long.reindex(np.random.permutation(df_long.index))
    #         df_long = df_long.reset_index(drop=True)
    #         if args['numlongquestions'] is not None:
    #             df_long = df_long.head(args['numlongquestions'])
    #
    #     doc = create_exam(df=df2, df_long=df_long, answer_columns=_columns, question=question, is_odd=i%2==0)
    #     file_name = f'{args["prefix"]}{i}.docx'
    #     doc.save(file_name)
    #     print(f'Saved file {file_name}')

    #    print_exam(df2, _columns, question)
# See PyCharm help at https://www.jetbrains.com/help/pycharm/
